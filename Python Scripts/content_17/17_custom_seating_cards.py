import os
import docx
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont

"""
Script 17_custom_seating_cards.py creates a party invitation for a list of guests.
A Word document generates with the predefined greeting and insertion of name.
A PNG file seating card is generated for each guest.
"""
def get_filepath(directory,target_filename):

    for folderpath,subfolder,filenames in os.walk(directory):

        for filename in filenames:

            if target_filename != filename:
                continue

            return os.path.join(folderpath,filename)

##Load invite File (Extract)
def import_invites(pDirectory,pFilename):

    """
    Loads invite file content into list invite_list

    Params:
    pDirectory: Folder location of file
    pFilename: Name of file

    Returns:
    invite_list: A list with all guest names
    """

    for folderpath,subfolder,filenames in os.walk(pDirectory):

        for filename in filenames:

            if filename != pFilename:
                continue

            invite_filepath = os.path.join(folderpath,filename)
            print(invite_filepath)

    invite_file = open(invite_filepath,'r')

    invite_list = invite_file.readlines()

    invite_file.close()

    return invite_list

##Transform word document
def add_invite(pDoc,pInviteList):

    """
    Creates invitation with pGuest name. All iterations save to the same word document.

    Params:
    pDoc: Word Document object.
    pInviteList: A list containing invite names

    Returns:
    pDoc: Word Document object with new guest invite.

    """

    for guest in pInviteList:

        paragraph = pDoc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run("It would be a pleasure to have the company of")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)

        paragraph = pDoc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run(guest.strip())
        run.bold = True
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)

        paragraph = pDoc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run("at")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)
        run.underline = True
        run = paragraph.add_run(" 1001 Test Street on the evening of ")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)

        paragraph = pDoc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run("April 1st")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)

        paragraph = pDoc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run("at")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)
        run.underline = True
        run = paragraph.add_run(" 7 o'clock")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)

        if "\n" in guest:
            pDoc.add_page_break()

    return pDoc

##TODO:Generate image file with name and logo

def generate_seating_card(invite_list):
    
    os.makedirs("seating cards",exist_ok=True)
    
    for invite in invite_list:
        x = 288
        y = 360
        filename = f"seating_card_{invite.strip()}.png"
        img = Image.new('RGB', (x,y), color = 'white')

        ## Add border
        draw = ImageDraw.Draw(img)
        draw.line([(0, 0), (x-1, 0), (x-1, y-1), (0, y-1), (0, 0)], fill='black',width=5)

        ## Add invite name

        try:
            font = ImageFont.truetype("arial.ttf", 40)
        except IOError:
            print("Font not found, using default font.")
            font = ImageFont.load_default()

        center_x = x//2
        center_y = y//2
        draw.text(
        (center_x, center_y),
        invite.strip(),
        font=font,
        fill='black',
        anchor="mm"
        )

        ## Add flower logo
        logo_filepath = get_filepath(os.getcwd(),"flower_logo.png")
        logoIm = Image.open(logo_filepath)
        logo_width, logo_height = int(x/4),int(y/4)
        logoIm = logoIm.resize((logo_width,logo_height))
        logoIm = logoIm.convert("RGBA")
        img = img.convert("RGBA")
        img.paste(logoIm, (x - logo_width -3, y - logo_height -3), logoIm)
        ## Save
        img.save(os.path.join("seating cards",filename), "PNG")


## Creates and exports invite word document
def main():

    """
    Calls all defined functions to generate a word document.
    """

    vDirectory = os.getcwd()
    vFilename = "13_invite.txt"

    invite_list = import_invites(vDirectory,vFilename)

    doc = docx.Document()

    doc = add_invite(doc,invite_list)

    doc.save("Guest_Invites.docx")

    generate_seating_card(invite_list)

    print("Script complete.")

main()
