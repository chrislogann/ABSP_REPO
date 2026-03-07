import os
import docx
from docx.shared import Pt

"""
Script 13_custom_invitations.py creates a party invitation for a list of guests.
A Word document generates with the predefined greeting and insertion of name.
"""

##Load invite File (Extract)
def import_invites(pDirectory,pFilename):

    """
    Loads invite file content into list invite_list

    Params:
    pDirectory: Folder location of file
    pFilename: Name of file

    Returns:
    invite_list: A list with all guest names
    """

    for folderpath,subfolder,filenames in os.walk(pDirectory):

        for filename in filenames:

            if filename != pFilename:
                continue

            invite_filepath = os.path.join(folderpath,filename)
            print(invite_filepath)

    invite_file = open(invite_filepath,'r')

    invite_list = invite_file.readlines()

    invite_file.close()

    return invite_list

##Transform word document
def add_invite(pDoc,pInviteList):

    """
    Creates invitation with pGuest name. All iterations save to the same word document.

    Params:
    pDoc: Word Document object.
    pInviteList: A list containing invite names

    Returns:
    pDoc: Word Document object with new guest invite.

    """

    for guest in pInviteList:

        paragraph = pDoc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run("It would be a pleasure to have the company of")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)

        paragraph = pDoc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run(guest.strip())
        run.bold = True
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)

        paragraph = pDoc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run("at")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)
        run.underline = True
        run = paragraph.add_run(" 1001 Test Street on the evening of ")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)

        paragraph = pDoc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run("April 1st")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)

        paragraph = pDoc.add_paragraph()
        paragraph.alignment = 1
        run = paragraph.add_run("at")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)
        run.underline = True
        run = paragraph.add_run(" 7 o'clock")
        font = run.font
        font.name = 'Comic Sans MS'
        font.size = Pt(14)

        if "\n" in guest:
            pDoc.add_page_break()

    return pDoc

## Creates and exports invite word document
def main():

    """
    Calls all defined functions to generate a word document.
    """

    vDirectory = os.getcwd()
    vFilename = "13_invite.txt"

    invite_list = import_invites(vDirectory,vFilename)

    doc = docx.Document()

    doc = add_invite(doc,invite_list)

    doc.save("Guest_Invites.docx")

    print("Script complete.")

main()
